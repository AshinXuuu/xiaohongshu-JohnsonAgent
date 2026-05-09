"""
读取 ../产品库/ 下所有品牌的 docx / pdf,提取文本写入 ../data/products.json。
业务同事新增产品后,在项目根目录运行:
    python scripts/build_products.py
然后 git commit + push 即可触发 Vercel 自动部署。
"""

import json
import sys
from pathlib import Path

try:
    from docx import Document
    from pypdf import PdfReader
except ImportError:
    print("请先安装依赖: pip install python-docx pypdf")
    sys.exit(1)

# 项目根目录(脚本位于 scripts/ 子目录)
ROOT = Path(__file__).resolve().parent.parent
PRODUCT_LIB = ROOT / "产品库"
OUTPUT = ROOT / "data" / "products.json"

# 品牌资料的 PDF 通常是 Brand Guidelines,以排版为主,
# 默认不放进 prompt(避免无效膨胀,也避免英文规范扰乱中文输出)。
# 如需放入,把下面这行改为 True
INCLUDE_BRAND_GUIDELINES = False

# 是否处理 PDF 文件。
# 默认 False:实际经验中 PDF 多为图片/扫描件,文字提取不稳定且大文件会卡住。
# 卖点提取主要依赖 docx,PDF 仅作为业务自查阅参考。
# 如果某个产品确实只有 PDF 没有 docx,且 PDF 是文本型(可复制文字),可以临时改为 True。
INCLUDE_PDF = False

# 跳过超过此大小的 PDF(单位 MB),避免大型扫描件卡死提取
PDF_MAX_MB = 5


def extract_docx(p: Path) -> str:
    try:
        doc = Document(p)
        parts = []
        for para in doc.paragraphs:
            t = para.text.strip()
            if t:
                parts.append(t)
        for table in doc.tables:
            for row in table.rows:
                cells = [c.text.strip() for c in row.cells if c.text.strip()]
                if cells:
                    parts.append(" | ".join(cells))
        return "\n".join(parts)
    except Exception as e:
        return f"[ERROR docx {p.name}: {e}]"


def extract_pdf(p: Path) -> str:
    try:
        reader = PdfReader(p)
        parts = []
        for i, page in enumerate(reader.pages):
            txt = page.extract_text()
            if txt and txt.strip():
                parts.append(f"--- Page {i+1} ---\n{txt.strip()}")
        return "\n\n".join(parts) if parts else ""
    except Exception as e:
        return f"[ERROR pdf {p.name}: {e}]"


def extract_dir(d: Path) -> str:
    """合并文件夹下所有 docx/pdf 文本(PDF 默认跳过,见顶部 INCLUDE_PDF 配置)"""
    chunks = []
    for f in sorted(d.iterdir()):
        if f.name.startswith("."):
            continue
        if f.suffix.lower() == ".docx":
            chunks.append(f"## 文件:{f.name}\n\n{extract_docx(f)}")
        elif f.suffix.lower() == ".pdf":
            if not INCLUDE_PDF:
                continue
            size_mb = f.stat().st_size / (1024 * 1024)
            if size_mb > PDF_MAX_MB:
                print(f"  ⏭️  跳过大 PDF: {f.name} ({size_mb:.1f}MB > {PDF_MAX_MB}MB)")
                continue
            print(f"  📄 提取 PDF: {f.name} ({size_mb:.1f}MB)")
            text = extract_pdf(f)
            if text:
                chunks.append(f"## 文件:{f.name}\n\n{text}")
    return "\n\n".join(chunks)


def main():
    if not PRODUCT_LIB.exists():
        print(f"错误:找不到产品库目录 {PRODUCT_LIB}")
        sys.exit(1)

    result = {"brands": []}

    for brand_dir in sorted(PRODUCT_LIB.iterdir()):
        if not brand_dir.is_dir() or brand_dir.name.startswith("."):
            continue
        brand_name = brand_dir.name
        brand = {"name": brand_name, "guidelines": "", "products": []}

        for sub in sorted(brand_dir.iterdir()):
            if not sub.is_dir() or sub.name.startswith("."):
                continue
            if sub.name == "品牌资料":
                if INCLUDE_BRAND_GUIDELINES:
                    brand["guidelines"] = extract_dir(sub)
            else:
                product_text = extract_dir(sub)
                brand["products"].append({
                    "name": sub.name,
                    "content": product_text or "(暂无产品资料,请在产品库相应文件夹中放入 docx/pdf 文件)"
                })

        result["brands"].append(brand)

    OUTPUT.parent.mkdir(parents=True, exist_ok=True)
    with OUTPUT.open("w", encoding="utf-8") as f:
        json.dump(result, f, ensure_ascii=False, indent=2)

    print(f"已生成 {OUTPUT}")
    for b in result["brands"]:
        print(f"  品牌: {b['name']}  品牌资料字数: {len(b['guidelines'])}")
        for p in b["products"]:
            print(f"    - {p['name']}: {len(p['content'])} 字符")


if __name__ == "__main__":
    main()
