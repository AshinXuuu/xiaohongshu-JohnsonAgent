"""批量 OCR 入库管道 — 扫 产品库/ 全量,豆包视觉 OCR + docx 抽取,落 SQLite。

特点:
  ✓ 断点续传 / 幂等:按 file_path + 文件 SHA256 hash 去重,已处理的跳过
  ✓ 4 路并发(可改 --workers),每页一个任务,负载均衡
  ✓ 实时进度条 + 累计成本(豆包视觉 ~0.005 元/页)
  ✓ 自动分类:文件名带「说明书 / OM / Owners」→ manual_pdf;
             带「单页」→ onepager_pdf;
             带「卖点」.docx → selling_docx
  ✓ 自动跳过英文版(文件名含 ENG / 英文)
  ✓ 失败重试 2 次,持久失败的产品记录到日志

用法:
    # 跑全量
    python3 scripts/ocr_ingest.py

    # 只跑某个品牌
    python3 scripts/ocr_ingest.py --brand 乔山Johnson

    # 干跑(只列出会处理什么,不真的调 API)
    python3 scripts/ocr_ingest.py --dry-run

    # 强制重跑某个文件(清掉已记录后再跑)
    python3 scripts/ocr_ingest.py --force --product TX3

数据库表结构:
    manuals       — 每条记录一段文字(PDF 一页 = 一条;docx 全文 = 一条)
    manuals_files — 已处理文件登记,用于幂等
"""
import os
import sys
import ssl
import json
import time
import base64
import hashlib
import sqlite3
import argparse
import threading
import urllib.request
import urllib.error
from pathlib import Path
from concurrent.futures import ThreadPoolExecutor, as_completed
from collections import defaultdict

import fitz  # PyMuPDF


ROOT = Path(__file__).resolve().parent.parent
DB_PATH = ROOT / 'data' / 'usage.db'
PRODUCTS_ROOT = ROOT / '产品库'
LOG_DIR = ROOT / 'out' / 'ocr_logs'

ARK_URL = 'https://ark.cn-beijing.volces.com/api/v3/chat/completions'
DOUBAO_VISION_MODEL = os.environ.get(
    'DOUBAO_VISION_MODEL', 'doubao-1-5-vision-pro-32k-250115'
)
RENDER_DPI = 200
HTTP_TIMEOUT = 90
COST_PER_PAGE_YUAN = 0.005  # 豆包视觉粗略单价,实际以账单为准

OCR_PROMPT = (
    "你是专业的 OCR 文字识别引擎。请逐字识别下面这张图(产品说明书的某一页)里的"
    "所有可见文字,严格按照原始版面顺序输出(从上到下、从左到右,分栏的图按栏顺序)。\n"
    "要求:\n"
    "1. 只输出识别到的文字本身,不要添加任何解释、概括、润色。\n"
    "2. 数字、单位(cm/kg/W/dB/CHP 等)、型号、英文,严格按原图照抄,不要纠正、不要补全。\n"
    "3. 段落之间空一行;同一段内换行用单 \\n。\n"
    "4. 如果图上有表格,用 Markdown 表格语法输出。\n"
    "5. 图上的纯图标/插画/示意图,不需要描述,直接跳过。\n"
    "6. 如果整页根本没有文字(例如纯封面图、纯产品照),只输出 '[无文字内容]'。"
)


# ──────────────────────────── 工具 ────────────────────────────

def _build_ssl_context():
    try:
        import certifi
        return ssl.create_default_context(cafile=certifi.where())
    except ImportError:
        return ssl.create_default_context()


SSL_CTX = _build_ssl_context()


def load_env():
    env_file = ROOT / '.env'
    if not env_file.exists():
        return
    for raw in env_file.read_text(encoding='utf-8-sig').splitlines():
        line = raw.strip().lstrip('﻿')
        if line.startswith('export '):
            line = line[7:].strip()
        if not line or line.startswith('#') or '=' not in line:
            continue
        k, v = line.split('=', 1)
        k, v = k.strip(), v.strip()
        if len(v) >= 2 and ((v[0] == v[-1] == '"') or (v[0] == v[-1] == "'")):
            v = v[1:-1]
        os.environ.setdefault(k, v)


def sha256_file(path: Path) -> str:
    h = hashlib.sha256()
    with path.open('rb') as f:
        for chunk in iter(lambda: f.read(65536), b''):
            h.update(chunk)
    return h.hexdigest()


def classify_file(path: Path) -> tuple[str, bool]:
    """返回 (source_type, is_english_skip)
    source_type: 'manual_pdf' / 'onepager_pdf' / 'selling_docx' / 'other_pdf'
    """
    name = path.name
    name_lower = name.lower()

    if name.endswith('.docx'):
        if '卖点' in name:
            return 'selling_docx', False
        return 'other_docx', False

    # PDF
    # 英文版判定:文件名含 ENG / 英文
    is_eng = ('ENG' in name) or ('英文' in name)
    if is_eng:
        return 'manual_pdf', True  # 跳过

    if '说明书' in name:
        return 'manual_pdf', False
    # OM (Owners Manual) / OG (Operation Guide) / QSM (Quick Start Manual) / 用户手册
    if any(k in name for k in ('OM', 'OG ', '_OG', 'Owners', 'owners', 'QSM', '用户手册', '使用手册')):
        return 'manual_pdf', False
    if '单页' in name or '画板' in name:
        return 'onepager_pdf', False
    # 产品详解 / 产品介绍 算作产品宣传册
    if '产品详解' in name or '产品介绍' in name:
        return 'onepager_pdf', False
    # 兜底:其它 PDF 当作单页/产品宣传册
    return 'onepager_pdf', False


# ────────────────────────── DB schema ──────────────────────────

def init_schema(conn: sqlite3.Connection):
    conn.executescript("""
        CREATE TABLE IF NOT EXISTS manuals (
            id           INTEGER PRIMARY KEY AUTOINCREMENT,
            brand        TEXT NOT NULL,
            product      TEXT NOT NULL,
            source_type  TEXT NOT NULL,
            source_file  TEXT NOT NULL,
            page_no      INTEGER,
            content      TEXT NOT NULL,
            char_count   INTEGER NOT NULL,
            created_at   INTEGER NOT NULL
        );
        CREATE INDEX IF NOT EXISTS idx_manuals_product ON manuals(brand, product);
        CREATE INDEX IF NOT EXISTS idx_manuals_source  ON manuals(source_file);

        CREATE TABLE IF NOT EXISTS manuals_files (
            file_path     TEXT PRIMARY KEY,
            file_hash     TEXT NOT NULL,
            brand         TEXT NOT NULL,
            product       TEXT NOT NULL,
            source_type   TEXT NOT NULL,
            total_pages   INTEGER,
            total_chars   INTEGER,
            cost_yuan     REAL,
            status        TEXT NOT NULL,
            error         TEXT,
            completed_at  INTEGER NOT NULL
        );
        CREATE INDEX IF NOT EXISTS idx_files_hash ON manuals_files(file_hash);
    """)
    conn.commit()


# ────────────────────────── 调用层 ──────────────────────────

def call_doubao_vision(image_b64: str, api_key: str, retries: int = 2):
    """单页 OCR,失败重试 2 次。返回 (text, err)"""
    payload = {
        'model': DOUBAO_VISION_MODEL,
        'messages': [{
            'role': 'user',
            'content': [
                {'type': 'text', 'text': OCR_PROMPT},
                {'type': 'image_url', 'image_url': {'url': f'data:image/png;base64,{image_b64}'}},
            ],
        }],
        'temperature': 0.0,
        'max_tokens': 4000,
    }
    last_err = None
    for attempt in range(retries + 1):
        try:
            req = urllib.request.Request(
                ARK_URL,
                data=json.dumps(payload).encode('utf-8'),
                headers={
                    'Content-Type': 'application/json',
                    'Authorization': f'Bearer {api_key}',
                },
                method='POST',
            )
            with urllib.request.urlopen(req, timeout=HTTP_TIMEOUT, context=SSL_CTX) as resp:
                data = json.loads(resp.read().decode('utf-8'))
            return data['choices'][0]['message']['content'], None
        except urllib.error.HTTPError as e:
            detail = e.read().decode('utf-8', errors='ignore') if e.fp else ''
            last_err = f'HTTP {e.code}: {detail[:300]}'
            # 4xx 别重试,5xx 才重试
            if 400 <= e.code < 500:
                return None, last_err
        except Exception as e:
            last_err = f'{type(e).__name__}: {e}'
        if attempt < retries:
            time.sleep(2 + attempt * 2)
    return None, last_err


def render_page_png(pdf_path: Path, page_idx: int, dpi: int = RENDER_DPI) -> bytes:
    doc = fitz.open(pdf_path)
    pix = doc[page_idx].get_pixmap(dpi=dpi)
    png = pix.tobytes('png')
    doc.close()
    return png


def extract_docx_text(path: Path) -> str:
    """抽 docx 全文(含表格)"""
    from docx import Document
    doc = Document(path)
    parts = [p.text for p in doc.paragraphs if p.text.strip()]
    for tbl in doc.tables:
        for row in tbl.rows:
            for cell in row.cells:
                t = cell.text.strip()
                if t:
                    parts.append(t)
    return '\n'.join(parts)


# ────────────────────────── 进度展示 ──────────────────────────

class Progress:
    """简单的多行进度展示,加锁防止并发输出错乱"""

    def __init__(self):
        self.lock = threading.Lock()
        self.total_cost = 0.0
        self.total_pages_done = 0
        self.start_time = time.time()

    def page_done(self, ok: bool):
        with self.lock:
            self.total_pages_done += 1
            if ok:
                self.total_cost += COST_PER_PAGE_YUAN

    def line(self, msg: str):
        with self.lock:
            elapsed = time.time() - self.start_time
            print(f'[{elapsed:>6.1f}s | 已花 {self.total_cost:>5.2f}元 | 已完成 {self.total_pages_done} 页] {msg}', flush=True)


# ────────────────────────── 文件处理 ──────────────────────────

def gather_files(brand_filter=None, product_filter=None):
    """扫 产品库/ 返回待处理的 (file_path, brand, product, source_type) 列表"""
    out = []
    for brand_dir in sorted(PRODUCTS_ROOT.iterdir()):
        if not brand_dir.is_dir() or brand_dir.name.startswith('.'):
            continue
        if brand_filter and brand_filter not in brand_dir.name:
            continue
        for prod_dir in sorted(brand_dir.iterdir()):
            if not prod_dir.is_dir() or prod_dir.name.startswith('.'):
                continue
            if product_filter and product_filter not in prod_dir.name:
                continue
            for f in sorted(prod_dir.iterdir()):
                if not f.is_file() or f.name.startswith('.'):
                    continue
                if f.suffix.lower() not in ('.pdf', '.docx'):
                    continue
                source_type, is_eng = classify_file(f)
                rel_path = str(f.relative_to(ROOT))
                out.append({
                    'rel_path': rel_path,
                    'abs_path': f,
                    'brand': brand_dir.name,
                    'product': prod_dir.name,
                    'source_type': source_type,
                    'is_eng_skip': is_eng,
                })
    return out


def already_processed(conn: sqlite3.Connection, file_path: str, file_hash: str) -> bool:
    row = conn.execute(
        "SELECT status FROM manuals_files WHERE file_path = ? AND file_hash = ?",
        (file_path, file_hash),
    ).fetchone()
    return row is not None and row[0] in ('success', 'skipped_english')


def clear_file_records(conn: sqlite3.Connection, file_path: str):
    """--force 时清掉旧记录"""
    row = conn.execute(
        "SELECT source_file FROM manuals_files WHERE file_path = ?",
        (file_path,),
    ).fetchone()
    if row:
        # 用文件名删 manuals 里的相关页(文件名应该全局唯一)
        conn.execute("DELETE FROM manuals WHERE source_file = ?", (Path(file_path).name,))
    conn.execute("DELETE FROM manuals_files WHERE file_path = ?", (file_path,))
    conn.commit()


def ocr_one_pdf(file_info, api_key, db_lock, conn_path, progress, workers_per_file=2):
    """对一个 PDF 跑全页 OCR,写入 SQLite。
    workers_per_file: 当前 PDF 内部并发数(配合外层多 PDF 并发,效率更高)
    """
    rel = file_info['rel_path']
    pdf_path = file_info['abs_path']
    brand = file_info['brand']
    product = file_info['product']
    source_type = file_info['source_type']
    source_file = pdf_path.name

    progress.line(f'▶ 开始 {brand}/{product} — {source_file}')

    try:
        doc = fitz.open(pdf_path)
        n = len(doc)
        doc.close()
    except Exception as e:
        return {'file': rel, 'ok': False, 'error': f'open failed: {e}', 'pages': 0, 'chars': 0}

    page_results = [None] * n
    page_errors = []

    def do_page(page_idx):
        try:
            # 先按默认 DPI 渲染;若图太大被豆包拒(InvalidParameter / Oversize)
            # 自动降 DPI 重试(200 → 150 → 100)
            for dpi in (RENDER_DPI, 150, 100):
                png = render_page_png(pdf_path, page_idx, dpi=dpi)
                if len(png) > 9 * 1024 * 1024 and dpi > 100:
                    # 图本身就已经 > 9MB 了,直接降一档
                    continue
                b64 = base64.b64encode(png).decode('ascii')
                text, err = call_doubao_vision(b64, api_key)
                if err and ('Oversize' in err or 'InvalidParameter' in err) and dpi > 100:
                    # 服务器侧拒了,降 DPI 重试
                    continue
                if err:
                    page_errors.append((page_idx + 1, err))
                    progress.page_done(ok=False)
                    return None
                page_results[page_idx] = text
                progress.page_done(ok=True)
                return text
            page_errors.append((page_idx + 1, 'oversize even at 100 DPI'))
            progress.page_done(ok=False)
            return None
        except Exception as e:
            page_errors.append((page_idx + 1, f'{type(e).__name__}: {e}'))
            progress.page_done(ok=False)
            return None

    with ThreadPoolExecutor(max_workers=workers_per_file) as ex:
        list(ex.map(do_page, range(n)))

    # 写入 SQLite
    total_chars = 0
    success_pages = 0
    with db_lock:
        conn = sqlite3.connect(str(conn_path))
        try:
            for idx, text in enumerate(page_results):
                if text is None:
                    continue
                conn.execute(
                    "INSERT INTO manuals (brand, product, source_type, source_file, "
                    "page_no, content, char_count, created_at) VALUES (?,?,?,?,?,?,?,?)",
                    (brand, product, source_type, source_file, idx + 1,
                     text, len(text), int(time.time() * 1000))
                )
                total_chars += len(text)
                success_pages += 1
            conn.commit()
        finally:
            conn.close()

    progress.line(
        f'✓ 完成 {source_file}{success_pages}/{n} 页,{total_chars} 字'
        + (f',{len(page_errors)} 页失败' if page_errors else '')
    )

    return {
        'file': rel,
        'ok': success_pages == n,
        'pages_done': success_pages,
        'total_pages': n,
        'chars': total_chars,
        'errors': page_errors,
        'cost': success_pages * COST_PER_PAGE_YUAN,
    }


def ingest_one_docx(file_info, db_lock, conn_path, progress):
    """docx 直接抽文本入库(零成本)"""
    rel = file_info['rel_path']
    docx_path = file_info['abs_path']
    brand = file_info['brand']
    product = file_info['product']
    source_type = file_info['source_type']
    source_file = docx_path.name

    try:
        text = extract_docx_text(docx_path)
    except Exception as e:
        return {'file': rel, 'ok': False, 'error': f'extract failed: {e}', 'chars': 0}

    if not text.strip():
        return {'file': rel, 'ok': False, 'error': 'empty content', 'chars': 0}

    with db_lock:
        conn = sqlite3.connect(str(conn_path))
        try:
            conn.execute(
                "INSERT INTO manuals (brand, product, source_type, source_file, "
                "page_no, content, char_count, created_at) VALUES (?,?,?,?,?,?,?,?)",
                (brand, product, source_type, source_file, 1,
                 text, len(text), int(time.time() * 1000))
            )
            conn.commit()
        finally:
            conn.close()

    progress.line(f'✓ docx {source_file} — {len(text)} 字')
    return {
        'file': rel,
        'ok': True,
        'pages_done': 1,
        'total_pages': 1,
        'chars': len(text),
        'errors': [],
        'cost': 0.0,
    }


def register_file(conn_path, db_lock, file_info, file_hash, result, status):
    with db_lock:
        conn = sqlite3.connect(str(conn_path))
        try:
            conn.execute(
                "INSERT OR REPLACE INTO manuals_files (file_path, file_hash, brand, product, "
                "source_type, total_pages, total_chars, cost_yuan, status, error, completed_at) "
                "VALUES (?,?,?,?,?,?,?,?,?,?,?)",
                (
                    file_info['rel_path'],
                    file_hash,
                    file_info['brand'],
                    file_info['product'],
                    file_info['source_type'],
                    result.get('total_pages'),
                    result.get('chars'),
                    result.get('cost', 0.0),
                    status,
                    result.get('error') or (
                        '; '.join(f'p{p}:{e[:50]}' for p, e in result.get('errors', [])[:3])
                        or None
                    ),
                    int(time.time() * 1000),
                )
            )
            conn.commit()
        finally:
            conn.close()


# ────────────────────────── main ──────────────────────────

def main():
    parser = argparse.ArgumentParser()
    parser.add_argument('--brand', help='只处理含此关键字的品牌')
    parser.add_argument('--product', help='只处理含此关键字的产品')
    parser.add_argument('--workers', type=int, default=4, help='文件级并发数,默认 4')
    parser.add_argument('--per-file-workers', type=int, default=2, help='单文件内部并发数,默认 2')
    parser.add_argument('--force', action='store_true', help='强制重跑,清掉旧记录')
    parser.add_argument('--dry-run', action='store_true', help='只列清单,不真正调 API')
    args = parser.parse_args()

    load_env()
    api_key = os.environ.get('DOUBAO_API_KEY', '').strip()
    if not api_key and not args.dry_run:
        print('❌ DOUBAO_API_KEY 未配置,请先在 .env 里填上')
        sys.exit(1)

    DB_PATH.parent.mkdir(parents=True, exist_ok=True)
    LOG_DIR.mkdir(parents=True, exist_ok=True)
    conn = sqlite3.connect(str(DB_PATH))
    init_schema(conn)
    conn.close()

    files = gather_files(args.brand, args.product)

    # 分类统计
    by_type = defaultdict(int)
    for f in files:
        by_type[f['source_type']] += 1
    print(f'\n扫到 {len(files)} 个文件:')
    for t, c in sorted(by_type.items(), key=lambda x: -x[1]):
        print(f'  {c:>3d} × {t}')
    print()

    # 计算每份的 hash + 检查已处理
    conn = sqlite3.connect(str(DB_PATH))
    pending_pdf = []
    pending_docx = []
    skipped_eng = 0
    skipped_existing = 0
    for f in files:
        f['file_hash'] = sha256_file(f['abs_path'])
        if f['is_eng_skip']:
            skipped_eng += 1
            register_file(DB_PATH, threading.Lock(), f, f['file_hash'], {}, 'skipped_english')
            continue
        if not args.force and already_processed(conn, f['rel_path'], f['file_hash']):
            skipped_existing += 1
            continue
        if args.force:
            clear_file_records(conn, f['rel_path'])
        if f['abs_path'].suffix.lower() == '.docx':
            pending_docx.append(f)
        else:
            pending_pdf.append(f)
    conn.close()

    # 预估成本和时间
    total_pdf_pages = 0
    for f in pending_pdf:
        try:
            d = fitz.open(f['abs_path'])
            total_pdf_pages += len(d)
            d.close()
        except Exception:
            pass

    print(f'跳过英文版: {skipped_eng}')
    print(f'跳过已处理: {skipped_existing}')
    print(f'待 OCR PDF: {len(pending_pdf)} 个文件 ≈ {total_pdf_pages} 页 ≈ {total_pdf_pages * COST_PER_PAGE_YUAN:.2f} 元')
    print(f'待抽 docx: {len(pending_docx)} 个文件(免费)')
    print(f'并发: {args.workers} 文件 × {args.per_file_workers} 页/文件')
    print()

    if args.dry_run:
        print('═══ DRY RUN 待处理清单 ═══')
        for f in pending_pdf:
            print(f"  PDF  {f['brand']} / {f['product']} / {f['abs_path'].name}")
        for f in pending_docx:
            print(f"  DOCX {f['brand']} / {f['product']} / {f['abs_path'].name}")
        return

    if not pending_pdf and not pending_docx:
        print('没有要处理的文件 — 全部已处理过或已跳过。如要重跑用 --force')
        return

    input('回车开始,Ctrl+C 取消 → ')

    db_lock = threading.Lock()
    progress = Progress()

    # docx 先快速串行处理(几秒搞定)
    for f in pending_docx:
        result = ingest_one_docx(f, db_lock, DB_PATH, progress)
        status = 'success' if result['ok'] else 'failed'
        register_file(DB_PATH, db_lock, f, f['file_hash'], result, status)

    # PDF 并发处理
    def process_pdf(f):
        return f, ocr_one_pdf(f, api_key, db_lock, DB_PATH, progress, args.per_file_workers)

    with ThreadPoolExecutor(max_workers=args.workers) as ex:
        futures = [ex.submit(process_pdf, f) for f in pending_pdf]
        for fut in as_completed(futures):
            try:
                f, result = fut.result()
                status = 'success' if result['ok'] else 'partial'
                register_file(DB_PATH, db_lock, f, f['file_hash'], result, status)
            except Exception as e:
                progress.line(f'✗ 任务异常: {e}')

    # 汇总
    print()
    print('═══════════════════════════════')
    print(f'全部完成,总耗时 {time.time() - progress.start_time:.1f}s')
    print(f'累计成本 ≈ {progress.total_cost:.2f} 元')
    print(f'累计 OCR 页数: {progress.total_pages_done}')
    print()
    print(f'查询入库情况:')
    print(f'  sqlite3 {DB_PATH} "SELECT brand, product, source_type, COUNT(*) FROM manuals GROUP BY 1,2,3 ORDER BY 1,2;"')


if __name__ == '__main__':
    main()
